"""Text extraction for PDF (digital/scanned) and DOCX.

Returns page-segmented text so downstream chunking can track page boundaries.
"""

import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from shared.text_normalize import clean_for_ingestion, strip_headers_footers

log = logging.getLogger("gutenberg.extractors")


def extract_modal_elements(path: Path, doc_type: str) -> list[dict]:
    """Return tables and equations as structured modal elements (P3).

    Best-effort via PyMuPDF for digital PDFs (tables via ``find_tables``).
    Empty list for scanned PDFs / DOCX / EPUB — upstream modal processors
    rely on Docling/Surya paths that aren't threaded through yet.
    """
    try:
        from core.config import settings
    except Exception:
        from services.api.core.config import settings  # type: ignore

    if not getattr(settings, "feature_modal_chunks", False):
        return []
    if doc_type not in {"pdf_digital", "pdf_scanned"}:
        return []

    elements: list[dict] = []
    try:
        doc = fitz.open(str(path))
        for i, page in enumerate(doc):
            try:
                tables = page.find_tables()
            except Exception:
                tables = []
            for t in tables or []:
                try:
                    md = t.to_markdown() if hasattr(t, "to_markdown") else ""
                except Exception:
                    md = ""
                if md.strip():
                    elements.append(
                        {
                            "kind": "table",
                            "content": md,
                            "page": i + 1,
                            "surrounding_text": page.get_text()[:1500],
                        }
                    )
        doc.close()
    except Exception as e:
        log.debug(f"modal extraction skipped for {path.name}: {e}")
    return elements


def extract_text(path: Path, doc_type: str, strategy: str = "default") -> tuple[str, dict, list[dict]]:
    """Extract text and metadata from a document.

    Returns (full_text, metadata_dict, page_segments).
    page_segments is a list of {"page": int, "text": str} dicts.
    For formats without page numbers, page_segments is empty.

    strategy: "default" or "optimized" (for scanned PDFs).

    If a Surya markdown file exists for this document (in a configured
    surya corpus directory), it is used preferentially over PyMuPDF/Docling
    because Surya produces cleaner text for scanned and mixed-layout PDFs.
    """
    metadata = {
        "source": path.name,
        "doc_type": doc_type,
    }

    # Check for pre-existing Surya output (higher quality than PyMuPDF for most PDFs)
    surya_result = _try_surya_extraction(path, metadata)
    if surya_result is not None:
        return surya_result

    if doc_type == "pdf_digital":
        return _extract_pdf_digital(path, metadata)
    elif doc_type == "pdf_scanned":
        return _extract_pdf_scanned(path, metadata, strategy=strategy)
    elif doc_type == "docx":
        return _extract_docx(path, metadata)
    elif doc_type == "epub":
        from pipeline.epub_extractor import extract_epub
        return extract_epub(path)
    else:
        raise ValueError(f"Unknown doc_type: {doc_type}")


def _try_surya_extraction(path: Path, metadata: dict) -> tuple[str, dict, list[dict]] | None:
    """Check for pre-existing Surya markdown output and use it if available.

    Surya produces layout-aware OCR with proper paragraph structure,
    outperforming PyMuPDF on scanned and mixed-layout PDFs. The markdown
    files are expected at {SURYA_CORPUS_DIR}/{stem}/{stem}.md.

    Returns None if no Surya output exists for this document.
    """
    import os
    import re

    surya_dir = Path(os.environ.get("SURYA_CORPUS_DIR", "data/surya_corpus"))
    if not surya_dir.exists():
        return None

    # Match by stem (filename without extension)
    stem = path.stem
    surya_book_dir = surya_dir / stem
    if not surya_book_dir.exists():
        return None

    # Find the markdown file
    md_files = list(surya_book_dir.glob("*.md"))
    if not md_files:
        return None

    md_path = md_files[0]
    text = md_path.read_text(encoding="utf-8")
    if not text.strip():
        return None

    log.info(f"Using Surya extraction for {path.name} ({len(text):,} chars)")

    # Clean: strip markdown image references
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", text)

    # Build page segments using PyMuPDF page count for rough mapping
    # Surya markdown doesn't have page markers, so we use the PDF for page count
    # and distribute text proportionally
    try:
        doc = fitz.open(str(path))
        total_pages = len(doc)
        doc.close()
    except Exception:
        total_pages = 1

    # Split on markdown headers as rough page/section boundaries
    text = clean_for_ingestion(text)
    pages = [{"page": 1, "text": text}]

    metadata["total_pages"] = total_pages
    metadata["extraction"] = "surya"
    return text, metadata, pages


def _extract_pdf_digital(path: Path, metadata: dict) -> tuple[str, dict, list[dict]]:
    """Fast extraction using PyMuPDF for digital PDFs."""
    doc = fitz.open(str(path))
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append({"page": i + 1, "text": clean_for_ingestion(text)})
    doc.close()

    pages = strip_headers_footers(pages)
    metadata["total_pages"] = len(pages)
    full_text = "\n\n".join(p["text"] for p in pages)
    return full_text, metadata, pages


def _extract_pdf_scanned(path: Path, metadata: dict, strategy: str = "default") -> tuple[str, dict, list[dict]]:
    """OCR extraction using Docling for scanned PDFs with per-page text.

    Uses GPU acceleration when available (CUDA) for layout detection and OCR.
    strategy: "default" (full Docling) or "optimized" (skip tables/pictures/formulas, batch processing).
    """
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.datamodel.accelerator_options import AcceleratorDevice, AcceleratorOptions

    import torch
    device = AcceleratorDevice.CUDA if torch.cuda.is_available() else AcceleratorDevice.CPU
    log.info(f"Docling using device: {device}, strategy: {strategy}")

    pipeline_options = PdfPipelineOptions(
        accelerator_options=AcceleratorOptions(device=device),
    )

    if strategy == "optimized":
        # Skip features unnecessary for philosophical books (no tables, figures, formulas)
        pipeline_options.do_table_structure = False
        pipeline_options.do_picture_classification = False
        pipeline_options.do_formula_enrichment = False
        pipeline_options.do_code_enrichment = False
        # Conservative batch sizes for 8GB GPU (layout model needs ~2GB + batch overhead)
        pipeline_options.ocr_batch_size = 8
        pipeline_options.layout_batch_size = 4
        log.info("Docling optimized: tables/pictures/formulas OFF, layout_batch=4, ocr_batch=8")

    converter = DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
        }
    )
    result = converter.convert(str(path))
    doc = result.document

    # Try per-page extraction via Docling's page-aware markdown export
    pages: list[dict] = []
    try:
        num_pages = len(doc.pages) if hasattr(doc, "pages") and doc.pages else 0
        if num_pages > 0:
            for page_no in sorted(doc.pages.keys()):
                page_text = doc.export_to_markdown(page_no=page_no)
                if page_text.strip():
                    pages.append({"page": page_no, "text": clean_for_ingestion(page_text)})
            log.info(f"Docling per-page extraction: {len(pages)} pages with text out of {num_pages}")
    except (TypeError, AttributeError) as e:
        log.warning(f"Docling per-page extraction failed, falling back to full export: {e}")
        pages = []

    # Fallback: single-page segment from full markdown export
    if not pages:
        md_text = clean_for_ingestion(doc.export_to_markdown())
        total = len(doc.pages) if hasattr(doc, "pages") and doc.pages else 0
        metadata["total_pages"] = total
        metadata["ocr"] = True
        page_segments = [{"page": 1, "text": md_text}] if md_text.strip() else []
        return md_text, metadata, page_segments

    pages = strip_headers_footers(pages)
    metadata["total_pages"] = len(pages)
    metadata["ocr"] = True

    # OCR quality heuristic: check for garbled text
    _check_ocr_quality(pages, metadata)

    full_text = "\n\n".join(p["text"] for p in pages)
    return full_text, metadata, pages


def _check_ocr_quality(pages: list[dict], metadata: dict):
    """Warn if OCR quality appears low using SpaCy POS tagging.

    Tokens tagged as X (unknown POS) by SpaCy are likely OCR garbage.
    Falls back to regex heuristic if SpaCy is unavailable.
    """
    sample_text = " ".join(p["text"] for p in pages[:5])[:10000]
    if not sample_text.strip():
        return

    try:
        from shared.nlp import get_nlp, is_available
        if is_available():
            nlp = get_nlp()
            doc = nlp(sample_text)
            total = sum(1 for t in doc if t.is_alpha)
            garbage = sum(1 for t in doc if t.is_alpha and t.pos_ == "X")
            ratio = garbage / total if total > 0 else 0
            if ratio > 0.15:
                log.warning(
                    f"Low OCR quality for {metadata.get('source', 'unknown')}: "
                    f"{ratio:.0%} unrecognized tokens (SpaCy POS=X)"
                )
                metadata["ocr_quality"] = "low"
            return
    except ImportError:
        pass

    # Fallback: regex heuristic
    import re
    words = re.findall(r"[a-zA-Z]{3,}", sample_text)
    if not words:
        return
    garbled = sum(1 for w in words if len(w) <= 2 or not w.isascii())
    ratio = garbled / len(words) if words else 0
    if ratio > 0.20:
        log.warning(
            f"Low OCR quality for {metadata.get('source', 'unknown')}: "
            f"{ratio:.0%} potentially garbled words"
        )


def _extract_docx(path: Path, metadata: dict) -> tuple[str, dict, list[dict]]:
    """Extract text from DOCX using python-docx."""
    doc = DocxDocument(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Preserve heading structure for chunker
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    metadata["total_pages"] = 0  # DOCX doesn't have page numbers
    full_text = clean_for_ingestion("\n\n".join(paragraphs))
    return full_text, metadata, []  # no page segments for DOCX
