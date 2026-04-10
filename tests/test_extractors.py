import sys
from pathlib import Path
from unittest.mock import MagicMock, patch

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "services" / "worker"))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "services"))

import pytest
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from pipeline.extractors import (
    _extract_pdf_digital,
    _check_ocr_quality,
    _extract_docx,
    extract_text,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def simple_pdf(tmp_path):
    """Create a minimal 2-page PDF with known text."""
    doc = fitz.open()
    for i, body in enumerate(["First page content here.", "Second page with more text."], 1):
        page = doc.new_page()
        page.insert_text((72, 72), body)
    pdf_path = tmp_path / "simple.pdf"
    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


@pytest.fixture
def simple_docx(tmp_path):
    """Create a minimal DOCX with headings and body paragraphs."""
    doc = DocxDocument()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_heading("Chapter One", level=2)
    doc.add_paragraph("Chapter one body text.")
    docx_path = tmp_path / "simple.docx"
    doc.save(str(docx_path))
    return docx_path


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestExtractPdfDigital:
    def test_returns_page_dicts(self, simple_pdf):
        """_extract_pdf_digital returns list of {page, text} dicts."""
        _full_text, _meta, pages = _extract_pdf_digital(
            simple_pdf, {"source": "simple.pdf", "doc_type": "pdf_digital"}
        )
        assert isinstance(pages, list)
        assert len(pages) == 2
        for p in pages:
            assert "page" in p and "text" in p
            assert isinstance(p["page"], int)
            assert isinstance(p["text"], str)
        assert pages[0]["page"] == 1
        assert pages[1]["page"] == 2
        assert "First page" in pages[0]["text"]
        assert "Second page" in pages[1]["text"]

    def test_metadata_has_total_pages(self, simple_pdf):
        """Metadata dict includes total_pages after extraction."""
        metadata = {"source": "simple.pdf", "doc_type": "pdf_digital"}
        _full_text, meta, pages = _extract_pdf_digital(simple_pdf, metadata)
        assert "total_pages" in meta
        assert meta["total_pages"] == len(pages)
        assert meta["source"] == "simple.pdf"


class TestCheckOcrQuality:
    def test_good_text_no_low_quality_flag(self):
        """Normal English text should not trigger low-quality warning."""
        pages = [
            {"page": 1, "text": "Philosophy is the study of general and fundamental questions."},
            {"page": 2, "text": "Knowledge, reality, and existence are central concerns."},
        ]
        metadata = {"source": "good.pdf"}

        # Mock shared.nlp to simulate spacy tagging all tokens as real words
        mock_nlp = MagicMock()
        mock_doc = []
        for word in "Philosophy is the study of general questions".split():
            tok = MagicMock()
            tok.is_alpha = True
            tok.pos_ = "NOUN"
            mock_doc.append(tok)
        mock_nlp.return_value = mock_doc

        with patch("pipeline.extractors.get_nlp", mock_nlp, create=True), \
             patch.dict("sys.modules", {"shared.nlp": MagicMock(
                 get_nlp=MagicMock(return_value=mock_nlp),
                 is_available=MagicMock(return_value=True),
             )}):
            _check_ocr_quality(pages, metadata)

        assert metadata.get("ocr_quality") != "low"

    def test_garbled_text_flags_low_quality(self):
        """Mostly unknown POS tokens should flag ocr_quality as low."""
        pages = [
            {"page": 1, "text": "xkzpq wvmbn trlgh dfkjx mnbvc"},
        ]
        metadata = {"source": "garbled.pdf"}

        # Build mock tokens: all tagged as X (unknown)
        mock_tokens = []
        for _ in range(20):
            tok = MagicMock()
            tok.is_alpha = True
            tok.pos_ = "X"
            mock_tokens.append(tok)

        mock_nlp_fn = MagicMock(return_value=mock_tokens)
        mock_nlp_module = MagicMock()
        mock_nlp_module.get_nlp.return_value = mock_nlp_fn
        mock_nlp_module.is_available.return_value = True

        with patch.dict("sys.modules", {"shared.nlp": mock_nlp_module}):
            _check_ocr_quality(pages, metadata)

        assert metadata.get("ocr_quality") == "low"


class TestExtractDocx:
    def test_preserves_headings_as_markdown(self, simple_docx):
        """Headings in DOCX should be converted to markdown # syntax."""
        metadata = {"source": "simple.docx", "doc_type": "docx"}
        full_text, meta, page_segments = _extract_docx(simple_docx, metadata)

        assert "# Introduction" in full_text
        assert "## Chapter One" in full_text
        assert "introduction paragraph" in full_text.lower()
        assert page_segments == []
        assert meta["total_pages"] == 0


class TestExtractText:
    def test_returns_correct_tuple_types(self, simple_pdf):
        """extract_text returns (str, dict, list) with correct types."""
        full_text, metadata, pages = extract_text(simple_pdf, "pdf_digital")

        assert isinstance(full_text, str)
        assert isinstance(metadata, dict)
        assert isinstance(pages, list)
        assert len(full_text) > 0
        assert metadata["doc_type"] == "pdf_digital"
        assert metadata["source"] == simple_pdf.name

    def test_unknown_doc_type_raises(self, simple_pdf):
        """Unknown doc_type should raise ValueError."""
        with pytest.raises(ValueError, match="Unknown doc_type"):
            extract_text(simple_pdf, "unknown_format")
